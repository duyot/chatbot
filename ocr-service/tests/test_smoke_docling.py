"""One real Docling inference run. Excluded from default runs by pytest.ini
(`-m "not slow"`) because it loads model weights.

Run explicitly inside the built image:
    docker run --rm --network none ocr-spike python -m pytest tests/test_smoke_docling.py -m slow -v
"""
import io

import pytest
from PIL import Image, ImageDraw


@pytest.mark.slow
def test_parse_bytes_extracts_text_from_a_rendered_image():
    from parser import parse_bytes

    img = Image.new("RGB", (900, 300), "white")
    draw = ImageDraw.Draw(img)
    draw.text((40, 60), "QUARTERLY REPORT", fill="black")
    draw.text((40, 160), "Revenue increased in the APAC region.", fill="black")
    buf = io.BytesIO()
    img.save(buf, format="PNG")

    body = parse_bytes(buf.getvalue(), filename="scan.png")

    assert body["schema_version"] == 1
    assert body["metadata"]["page_count"] >= 1
    assert body["elements"], "expected at least one element from the rendered text"
    for el in body["elements"]:
        assert el["type"] in {
            "heading", "paragraph", "list_item", "table",
            "caption", "figure", "page_header", "page_footer",
        }
        if el["bbox"] is not None:
            assert all(0.0 <= v <= 1.0 for v in el["bbox"])


@pytest.mark.slow
def test_parse_bytes_bbox_is_not_vertically_mirrored():
    """Guards the bottom-left -> top-left bbox flip specifically.

    `0.0 <= v <= 1.0` (the assertion above) can't catch a mirrored bbox — a
    flip bug still produces in-range numbers, just for the wrong half of the
    page. These rects get overlaid on page images in the UI, where a mirror
    would be silently wrong on screen but invisible to that assertion.

    The image has text confined to a single line near the top, with the
    entire lower half blank, so there is only one place a genuine detection
    can land: a correct top-left-origin bbox must have y0 (and y1) in the
    upper half of the page. A bottom-left origin left unconverted (or a
    flip applied twice) would report this same text in the lower half
    instead.
    """
    from parser import parse_bytes

    img = Image.new("RGB", (900, 600), "white")
    draw = ImageDraw.Draw(img)
    draw.text((40, 40), "TOP OF PAGE ONLY", fill="black")
    buf = io.BytesIO()
    img.save(buf, format="PNG")

    body = parse_bytes(buf.getvalue(), filename="top_text.png")

    elements_with_bbox = [el for el in body["elements"] if el["bbox"] is not None]
    assert elements_with_bbox, "expected at least one element with a bbox"
    for el in elements_with_bbox:
        y0, y1 = el["bbox"][1], el["bbox"][3]
        assert y0 < 0.5 and y1 < 0.5, (
            f"element {el['text']!r} bbox {el['bbox']} lands in the lower half of "
            "the page; the top-of-page text should never normalize below y=0.5 "
            "unless the bottom-left/top-left flip is wrong"
        )


@pytest.mark.slow
def test_parse_bytes_extracts_elements_from_a_real_docx():
    """FIX 1 (Critical): DOCX parsing was returning zero elements.

    Docling's Word backend emits no provenance and no pages at all
    (`dl_doc.pages == {}`, `result.pages == []`, every item's `prov == []`).
    The old code skipped every item whose `prov` was falsy, so `elements`
    and `pages` both came back empty for every DOCX — a mocked wire body
    can't catch this because the bug is entirely in how parser.py drives
    real Docling output for this input format. This test builds a real
    .docx (via python-docx, a transitive Docling dependency for its Word
    backend, so it is already in this image) with a heading, a paragraph,
    and a table, and asserts elements actually come out — with `bbox is
    None`, since DOCX carries no page geometry to normalize against.
    """
    import docx
    from parser import parse_bytes

    document = docx.Document()
    document.add_heading("Quarterly Report", level=1)
    document.add_paragraph("Revenue increased in the APAC region this quarter.")
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Region"
    table.rows[0].cells[1].text = "Revenue"
    table.rows[1].cells[0].text = "APAC"
    table.rows[1].cells[1].text = "1200"

    import io
    buf = io.BytesIO()
    document.save(buf)

    body = parse_bytes(buf.getvalue(), filename="report.docx")

    assert body["elements"], "expected non-empty elements for a real DOCX"
    assert body["pages"], "expected a synthesized page for a DOCX with no native pagination"

    types = [el["type"] for el in body["elements"]]
    assert "heading" in types, f"expected a heading element, got types={types}"
    assert "table" in types, f"expected a table element, got types={types}"

    heading_texts = [el["text"] for el in body["elements"] if el["type"] == "heading"]
    assert any("Quarterly Report" in t for t in heading_texts)

    table_texts = [el["text"] for el in body["elements"] if el["type"] == "table"]
    assert any("APAC" in t for t in table_texts)

    # DOCX carries no page geometry to normalize bboxes against.
    for el in body["elements"]:
        assert el["bbox"] is None, f"expected null bbox for DOCX element, got {el['bbox']!r}"

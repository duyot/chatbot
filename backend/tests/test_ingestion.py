import os
import pytest
from unittest.mock import MagicMock, patch

def test_chunk_text_splits_long_content():
    from app.services.ingestion import chunk_text
    text = "word " * 500  # 2500 chars
    parents, children_per_parent = chunk_text(text)
    assert len(parents) >= 1
    assert len(children_per_parent) == len(parents)
    all_children = [c for sub in children_per_parent for c in sub]
    assert len(all_children) > 1

def test_chunk_text_short_content_stays_one_chunk():
    from app.services.ingestion import chunk_text
    text = "Short paragraph."
    parents, children_per_parent = chunk_text(text)
    assert len(parents) == 1
    assert parents[0] == "Short paragraph."
    assert len(children_per_parent) == 1
    assert children_per_parent[0][0] == "Short paragraph."

def test_parse_document_docx_returns_native_page(tmp_path, mocker):
    """Legacy path (docling_enabled=False)."""
    from docx import Document as DocxDocument
    from app.services import ingestion
    from app.services.ingestion import parse_document
    mocker.patch.object(ingestion.settings, "docling_enabled", False)
    docx_path = tmp_path / "test.docx"
    doc = DocxDocument()
    doc.add_paragraph("Hello from DOCX")
    doc.save(str(docx_path))
    parsed = parse_document(str(docx_path), "test.docx")
    assert len(parsed.pages) == 1
    assert parsed.pages[0].source == "native"
    assert "Hello from DOCX" in parsed.pages[0].text
    assert parsed.metadata["mime_type"].endswith("wordprocessingml.document")


def test_parse_document_image_uses_ocr(tmp_path, mocker):
    """Legacy path (docling_enabled=False)."""
    from app.services import ingestion
    from app.services.ingestion import parse_document
    mocker.patch.object(ingestion.settings, "docling_enabled", False)
    img_path = tmp_path / "photo.png"
    img_path.write_bytes(b"\x89PNG fake bytes")
    mocker.patch.object(ingestion, "ocr_image_lines", return_value={
        "lines": [
            {"text": "scanned text", "bbox": [[0, 0], [100, 0], [100, 20], [0, 20]], "confidence": 0.92},
        ],
        "width": 200, "height": 100,
    })
    parsed = parse_document(str(img_path), "photo.png")
    assert len(parsed.pages) == 1
    assert parsed.pages[0].source == "ocr"
    assert parsed.pages[0].text == "scanned text"
    assert abs(parsed.pages[0].ocr_confidence - 0.92) < 1e-6
    assert parsed.metadata["ocr_engine"] == "paddleocr"


def test_parse_document_pdf_native_text_skips_ocr(tmp_path, mocker):
    """Legacy path (docling_enabled=False)."""
    import fitz
    from app.services import ingestion
    from app.services.ingestion import parse_document
    mocker.patch.object(ingestion.settings, "docling_enabled", False)
    pdf_path = tmp_path / "doc.pdf"
    d = fitz.open()
    page = d.new_page()
    page.insert_text((72, 72), "This is a native digital PDF with plenty of text content.")
    d.save(str(pdf_path))
    d.close()
    ocr_spy = mocker.patch.object(ingestion, "ocr_image_lines", return_value={"lines": [], "width": 0, "height": 0})
    parsed = parse_document(str(pdf_path), "doc.pdf")
    assert parsed.pages[0].source == "native"
    assert "native digital PDF" in parsed.pages[0].text
    ocr_spy.assert_not_called()


def test_parse_document_pdf_scanned_triggers_ocr(tmp_path, mocker):
    """Legacy path (docling_enabled=False)."""
    import fitz
    from app.services import ingestion
    from app.services.ingestion import parse_document
    mocker.patch.object(ingestion.settings, "docling_enabled", False)
    pdf_path = tmp_path / "scan.pdf"
    d = fitz.open()
    d.new_page()  # blank page: no text layer
    d.save(str(pdf_path))
    d.close()
    mocker.patch.object(ingestion, "ocr_image_lines", return_value={
        "lines": [
            {"text": "ocr extracted text", "bbox": [[0, 0], [50, 0], [50, 10], [0, 10]], "confidence": 0.7},
        ],
        "width": 100, "height": 100,
    })
    parsed = parse_document(str(pdf_path), "scan.pdf")
    assert parsed.pages[0].source == "ocr"
    assert parsed.pages[0].text == "ocr extracted text"
    assert parsed.metadata["ocr_pages"] == 1

def test_embed_chunks_calls_openai_and_returns_vectors():
    from app.services.ingestion import embed_chunks
    fake_embedding = [0.1] * 1536
    mock_item = MagicMock()
    mock_item.embedding = fake_embedding
    mock_response = MagicMock()
    mock_response.data = [mock_item]
    with patch("app.services.ingestion.OpenAI") as MockOpenAI:
        MockOpenAI.return_value.embeddings.create.return_value = mock_response
        result = embed_chunks(["some text"])
    assert len(result) == 1
    assert len(result[0]) == 1536

def test_store_chunks_inserts_rows():
    from app.services.ingestion import store_chunks, ParentChunk, ChildChunk
    from app.models import DocumentChunk
    import uuid

    mock_db = MagicMock()
    # flush must assign .id on parent rows; simulate by making add_all set ids via side_effect
    def add_all_side_effect(rows):
        for i, row in enumerate(rows):
            row.id = uuid.uuid4()
    mock_db.add_all.side_effect = add_all_side_effect

    doc_id = str(uuid.uuid4())
    parents = [ParentChunk(content="parent one", page_start=2, page_end=2, source="ocr")]
    children_per_parent = [[
        ChildChunk(content="chunk one", page=2, source="ocr", ocr_confidence=0.8),
        ChildChunk(content="chunk two", page=2, source="ocr", ocr_confidence=0.8),
    ]]
    embeddings = [[0.1] * 1536, [0.2] * 1536]

    store_chunks(mock_db, doc_id, parents, children_per_parent, embeddings)

    mock_db.add_all.assert_called_once()
    parent_objs = mock_db.add_all.call_args[0][0]
    assert parent_objs[0].page_start == 2
    assert parent_objs[0].source == "ocr"
    mock_db.flush.assert_called_once()
    mock_db.bulk_save_objects.assert_called_once()
    saved_objects = mock_db.bulk_save_objects.call_args[0][0]
    assert len(saved_objects) == 2
    assert isinstance(saved_objects[0], DocumentChunk)
    assert saved_objects[0].chunk_index == 0
    assert saved_objects[0].page == 2
    assert saved_objects[0].source == "ocr"
    assert saved_objects[1].chunk_index == 1
    mock_db.commit.assert_called_once()

def test_embed_text_calls_openai_and_returns_vector(mocker):
    mock_item = mocker.MagicMock()
    mock_item.embedding = [0.1, 0.2, 0.3]
    mock_response = mocker.MagicMock()
    mock_response.data = [mock_item]
    mock_client = mocker.MagicMock()
    mock_client.embeddings.create.return_value = mock_response
    mocker.patch("app.services.ingestion.OpenAI", return_value=mock_client)

    from app.services.ingestion import embed_text
    result = embed_text("hello world")

    assert result == [0.1, 0.2, 0.3]


def test_document_parent_chunk_model_round_trips(db):
    from app.models import Document, DocumentParentChunk, DocumentChunk
    import uuid

    doc = Document(id=uuid.uuid4(), file_name="t.pdf", file_path="/tmp/t.pdf", status="done")
    db.add(doc)
    parent = DocumentParentChunk(
        document_id=doc.id,
        parent_index=0,
        content="Parent body of text...",
    )
    db.add(parent)
    db.flush()

    child = DocumentChunk(
        document_id=doc.id,
        parent_id=parent.id,
        chunk_index=0,
        content="Child snippet",
        embedding=[0.0] * 1536,
    )
    db.add(child)
    db.flush()

    fetched = db.query(DocumentChunk).filter_by(id=child.id).one()
    assert fetched.parent_id == parent.id


def test_chunk_document_carries_page_and_source():
    from app.services.ingestion import chunk_document, ParsedDocument, PageContent
    parsed = ParsedDocument(
        pages=[
            PageContent(page=1, text="Native page text. " * 5, source="native"),
            PageContent(page=2, text="Scanned page text. " * 5, source="ocr", ocr_confidence=0.75),
        ],
        metadata={},
    )
    parents, children_per_parent = chunk_document(parsed)
    assert len(parents) >= 2
    page2_parents = [p for p in parents if p.page_start == 2]
    assert page2_parents and page2_parents[0].source == "ocr"
    for parent, children in zip(parents, children_per_parent):
        for c in children:
            assert c.page == parent.page_start
            assert c.source == parent.source
            if c.source == "ocr":
                assert abs(c.ocr_confidence - 0.75) < 1e-6


def test_chunk_document_skips_blank_pages():
    from app.services.ingestion import chunk_document, ParsedDocument, PageContent
    parsed = ParsedDocument(
        pages=[
            PageContent(page=1, text="   ", source="ocr", ocr_confidence=None),
            PageContent(page=2, text="Real content here.", source="native"),
        ],
        metadata={},
    )
    parents, children_per_parent = chunk_document(parsed)
    assert parents and all(p.page_start == 2 for p in parents)


def test_chunk_metadata_columns_round_trip(db):
    """Document/parent/child carry the new OCR + page metadata columns."""
    from app.models import Document, DocumentParentChunk, DocumentChunk
    import uuid

    doc = Document(
        id=uuid.uuid4(),
        file_name="scan.pdf",
        file_path="/tmp/scan.pdf",
        status="done",
        mime_type="application/pdf",
        page_count=3,
        doc_metadata={"ocr_engine": "paddleocr", "ocr_pages": 2, "native_pages": 1},
    )
    db.add(doc)
    parent = DocumentParentChunk(
        document_id=doc.id,
        parent_index=0,
        content="Parent on page 2",
        page_start=2,
        page_end=2,
        source="ocr",
    )
    db.add(parent)
    db.flush()

    child = DocumentChunk(
        document_id=doc.id,
        parent_id=parent.id,
        chunk_index=0,
        content="Child snippet",
        embedding=[0.0] * 1536,
        page=2,
        source="ocr",
        ocr_confidence=0.87,
    )
    db.add(child)
    db.flush()

    fetched_doc = db.query(Document).filter_by(id=doc.id).one()
    fetched_parent = db.query(DocumentParentChunk).filter_by(id=parent.id).one()
    fetched_child = db.query(DocumentChunk).filter_by(id=child.id).one()

    assert fetched_doc.mime_type == "application/pdf"
    assert fetched_doc.page_count == 3
    assert fetched_doc.doc_metadata["ocr_engine"] == "paddleocr"
    assert fetched_parent.page_start == 2
    assert fetched_parent.page_end == 2
    assert fetched_parent.source == "ocr"
    assert fetched_child.page == 2
    assert fetched_child.source == "ocr"
    assert abs(fetched_child.ocr_confidence - 0.87) < 1e-6


def test_quad_to_norm_rect_normalizes_and_clamps():
    from app.services.ingestion import _quad_to_norm_rect
    rect = _quad_to_norm_rect([[10, 20], [110, 20], [110, 60], [10, 60]], 200, 100)
    assert rect == [0.05, 0.2, 0.55, 0.6]
    # missing quad or dims -> None (unmappable)
    assert _quad_to_norm_rect(None, 200, 100) is None
    assert _quad_to_norm_rect([[0, 0]], 0, 0) is None


def test_line_spans_and_rects_for_span():
    from app.services.ingestion import LayoutLine, _line_spans, _rects_for_span
    lines = [
        LayoutLine(text="hello", bbox=[0.0, 0.0, 0.5, 0.1]),   # chars 0..5
        LayoutLine(text="world", bbox=[0.0, 0.2, 0.5, 0.3]),   # chars 6..11 (after "\n")
    ]
    spans = _line_spans(lines)
    assert [(s, e) for s, e, _ in spans] == [(0, 5), (6, 11)]
    # span within the first line only
    assert _rects_for_span(spans, 0, 5) == [[0.0, 0.0, 0.5, 0.1]]
    # span straddling both lines
    assert _rects_for_span(spans, 3, 8) == [[0.0, 0.0, 0.5, 0.1], [0.0, 0.2, 0.5, 0.3]]
    # lines with an empty bbox are skipped
    spans2 = _line_spans([LayoutLine(text="x", bbox=[])])
    assert _rects_for_span(spans2, 0, 1) == []


def test_chunk_document_attaches_bbox_from_page_lines():
    from app.services.ingestion import chunk_document, ParsedDocument, PageContent, LayoutLine
    page = PageContent(
        page=1,
        text="hello\nworld",
        source="ocr",
        ocr_confidence=0.9,
        lines=[
            LayoutLine(text="hello", bbox=[0.0, 0.0, 0.5, 0.1]),
            LayoutLine(text="world", bbox=[0.0, 0.2, 0.5, 0.3]),
        ],
    )
    parsed = ParsedDocument(pages=[page], metadata={})
    _parents, children_per_parent = chunk_document(parsed)
    child = children_per_parent[0][0]  # short text -> single child spanning both lines
    assert child.bbox == [[0.0, 0.0, 0.5, 0.1], [0.0, 0.2, 0.5, 0.3]]


def test_chunk_document_no_lines_yields_empty_bbox():
    from app.services.ingestion import chunk_document, ParsedDocument, PageContent
    parsed = ParsedDocument(
        pages=[PageContent(page=1, text="Some text without geometry.", source="native")],
        metadata={},
    )
    _parents, children_per_parent = chunk_document(parsed)
    assert children_per_parent[0][0].bbox == []


def test_chunk_text_produces_parents_and_children():
    from app.services.ingestion import chunk_text

    # Produce text large enough to split into multiple parents
    text = ("This is a sentence. " * 800)  # ~16k characters -> multiple 1500-token parents
    parents, children_by_parent = chunk_text(text)

    assert len(parents) >= 2, "should split into multiple parents"
    # children_by_parent is a list aligned with parents — each entry is a list of child strings
    assert len(children_by_parent) == len(parents)
    for parent_text, children in zip(parents, children_by_parent):
        assert children, "every parent must have at least one child"
        # children should be shorter than the parent
        joined = " ".join(children)
        # Tokens are smaller than chars, but length sanity: children re-joined should
        # roughly cover the parent (allow slack for overlap and whitespace)
        assert len(joined) >= len(parent_text) * 0.7


def test_search_text_is_generated_from_context_and_content(db):
    """search_text is a Postgres STORED generated column; it must concatenate
    context and content, and stay correct when context is NULL."""
    import uuid
    from app.models import Document, DocumentChunk

    doc = Document(id=uuid.uuid4(), file_name="t.pdf", file_path="/tmp/t.pdf", status="done")
    db.add(doc)
    with_ctx = DocumentChunk(
        document_id=doc.id, chunk_index=0, content="the rate is 40%",
        context="Section 3 of the lease agreement.", embedding=[0.0] * 1536,
    )
    without_ctx = DocumentChunk(
        document_id=doc.id, chunk_index=1, content="bare chunk",
        embedding=[0.0] * 1536,
    )
    db.add_all([with_ctx, without_ctx])
    db.flush()
    db.refresh(with_ctx)
    db.refresh(without_ctx)

    assert with_ctx.search_text == "Section 3 of the lease agreement. the rate is 40%"
    # NULL context must not null out the whole column
    assert without_ctx.search_text == " bare chunk"


def test_build_embedding_input_prefixes_context_when_present():
    from app.services.ingestion import build_embedding_input
    assert build_embedding_input("Section 3.", "the rate is 40%") == (
        "Section 3.\n\nthe rate is 40%"
    )


def test_build_embedding_input_returns_bare_content_when_no_context():
    """Disabling contextual embeddings must be a true no-op, so the no-context
    path has to stay byte-identical to the pre-feature behaviour."""
    from app.services.ingestion import build_embedding_input
    assert build_embedding_input(None, "bare") == "bare"
    assert build_embedding_input("", "bare") == "bare"


def test_store_chunks_persists_context(db):
    import uuid
    from app.models import Document, DocumentChunk
    from app.services.ingestion import ParentChunk, ChildChunk, store_chunks

    doc = Document(id=uuid.uuid4(), file_name="t.pdf", file_path="/tmp/t.pdf", status="done")
    db.add(doc)
    db.flush()

    parents = [ParentChunk(content="P", page_start=1, page_end=1, source="native")]
    children = [[
        ChildChunk(content="c1", page=1, source="native", context="CTX ONE"),
        ChildChunk(content="c2", page=1, source="native", context=None),
    ]]
    store_chunks(db, str(doc.id), parents, children, [[0.0] * 1536, [0.0] * 1536])

    rows = (
        db.query(DocumentChunk)
        .filter(DocumentChunk.document_id == doc.id)
        .order_by(DocumentChunk.chunk_index)
        .all()
    )
    assert [r.context for r in rows] == ["CTX ONE", None]


def _wire(elements, pages=None):
    return {
        "schema_version": 1,
        "metadata": {"page_count": len(pages or [1]), "ocr_pages": 1, "native_pages": 0},
        "pages": pages or [{"page": 1, "width": 600.0, "height": 800.0,
                            "source": "ocr", "ocr_confidence": 0.91}],
        "elements": elements,
    }


def test_parse_document_uses_remote_parser_and_keeps_elements(tmp_path, mocker):
    from app.services import ingestion
    pdf = tmp_path / "scan.pdf"
    pdf.write_bytes(b"%PDF-fake")
    mocker.patch.object(ingestion.settings, "docling_enabled", True)
    mocker.patch.object(ingestion, "parse_document_remote", return_value=_wire([
        {"id": "e0", "page": 1, "type": "heading", "level": 2,
         "text": "3.2 Revenue", "bbox": [0.1, 0.1, 0.6, 0.14], "confidence": None},
        {"id": "e1", "page": 1, "type": "table",
         "text": "| R | Q1 |\n|---|---|\n| APAC | 12 |",
         "bbox": [0.1, 0.2, 0.9, 0.5], "confidence": 0.9},
    ]))

    parsed = ingestion.parse_document(str(pdf), "scan.pdf")

    assert [e.type for e in parsed.elements] == ["heading", "table"]
    assert parsed.elements[0].level == 2
    assert parsed.elements[1].bbox == [0.1, 0.2, 0.9, 0.5]


def test_parse_document_builds_pages_from_elements(tmp_path, mocker):
    """contextualizer needs parsed.text and per-page text, so pages must still
    be populated even though chunking now works off elements."""
    from app.services import ingestion
    pdf = tmp_path / "scan.pdf"
    pdf.write_bytes(b"%PDF-fake")
    mocker.patch.object(ingestion.settings, "docling_enabled", True)
    mocker.patch.object(ingestion, "parse_document_remote", return_value=_wire(
        [
            {"id": "e0", "page": 1, "type": "paragraph", "text": "page one prose",
             "bbox": None, "confidence": None},
            {"id": "e1", "page": 2, "type": "paragraph", "text": "page two prose",
             "bbox": None, "confidence": None},
        ],
        pages=[
            {"page": 1, "width": 600.0, "height": 800.0, "source": "ocr",
             "ocr_confidence": 0.9},
            {"page": 2, "width": 600.0, "height": 800.0, "source": "native",
             "ocr_confidence": None},
        ],
    ))

    parsed = ingestion.parse_document(str(pdf), "scan.pdf")

    assert [p.page for p in parsed.pages] == [1, 2]
    assert parsed.pages[0].text == "page one prose"
    assert parsed.pages[0].source == "ocr"
    assert parsed.pages[1].source == "native"
    assert "page one prose" in parsed.text and "page two prose" in parsed.text


def test_parse_document_sets_mime_type_and_engine(tmp_path, mocker):
    from app.services import ingestion
    pdf = tmp_path / "scan.pdf"
    pdf.write_bytes(b"%PDF-fake")
    mocker.patch.object(ingestion.settings, "docling_enabled", True)
    mocker.patch.object(ingestion, "parse_document_remote", return_value=_wire([]))

    parsed = ingestion.parse_document(str(pdf), "scan.pdf")

    assert parsed.metadata["mime_type"] == "application/pdf"
    assert parsed.metadata["engine"] == "docling"


def test_parse_document_legacy_path_when_docling_disabled(tmp_path, mocker):
    """docling_enabled=False must reproduce the old behaviour exactly."""
    from docx import Document as DocxDocument
    from app.services import ingestion
    docx_path = tmp_path / "legacy.docx"
    doc = DocxDocument()
    doc.add_paragraph("Legacy path still works")
    doc.save(str(docx_path))
    mocker.patch.object(ingestion.settings, "docling_enabled", False)
    remote = mocker.patch.object(ingestion, "parse_document_remote")

    parsed = ingestion.parse_document(str(docx_path), "legacy.docx")

    remote.assert_not_called()
    assert parsed.elements == []
    assert "Legacy path still works" in parsed.pages[0].text


def test_parse_document_warns_on_element_page_mismatch(tmp_path, mocker, caplog):
    """An element referencing a page absent from the wire body's `pages` list
    must not silently vanish from parsed.text — it should log a warning, and
    the well-formed pages must still come through intact."""
    import logging
    from app.services import ingestion
    pdf = tmp_path / "scan.pdf"
    pdf.write_bytes(b"%PDF-fake")
    mocker.patch.object(ingestion.settings, "docling_enabled", True)
    mocker.patch.object(ingestion, "parse_document_remote", return_value=_wire(
        [
            {"id": "e0", "page": 1, "type": "paragraph", "text": "page one prose",
             "bbox": None, "confidence": None},
            {"id": "e1", "page": 99, "type": "paragraph", "text": "orphaned prose",
             "bbox": None, "confidence": None},
        ],
        pages=[
            {"page": 1, "width": 600.0, "height": 800.0, "source": "ocr",
             "ocr_confidence": 0.9},
        ],
    ))

    with caplog.at_level(logging.WARNING, logger="app.services.ingestion"):
        parsed = ingestion.parse_document(str(pdf), "scan.pdf")

    assert [p.page for p in parsed.pages] == [1]
    assert parsed.pages[0].text == "page one prose"
    assert "orphaned prose" not in parsed.text
    warnings = [r for r in caplog.records if r.levelno == logging.WARNING]
    assert len(warnings) == 1
    assert "99" in warnings[0].getMessage()

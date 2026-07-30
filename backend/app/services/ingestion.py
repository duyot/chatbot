import logging
import os
from dataclasses import dataclass, field
from typing import List, Optional, Tuple
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from openai import OpenAI
from sqlalchemy.orm import Session

from ..config import settings
from ..models import DocumentChunk, DocumentParentChunk
from .ocr_client import ocr_image_lines, parse_document_remote, OCRError

logger = logging.getLogger(__name__)


# --- Parsed representation --------------------------------------------------

@dataclass
class LayoutLine:
    """One text line of a page with its axis-aligned bbox, normalized to [0,1]
    of the page (so it is independent of render DPI / display size)."""
    text: str
    bbox: List[float]                  # [x0, y0, x1, y1] in [0,1]


@dataclass
class PageContent:
    """One logical page of a source document.

    `text` is the page text used for chunking/embedding; when `lines` is present
    `text` is exactly "\\n".join(l.text for l in lines) so a chunk's character
    offset in `text` maps back to the contributing lines' bboxes.
    """
    page: int                          # 1-based
    text: str
    source: str                        # "native" | "ocr"
    ocr_confidence: Optional[float] = None
    lines: List[LayoutLine] = field(default_factory=list)
    width: Optional[float] = None      # source page width (px or pt), pre-normalization
    height: Optional[float] = None


@dataclass
class Element:
    """One typed document element from ocr-service /parse.

    `bbox` is [x0, y0, x1, y1] normalized to [0,1] of its page, or None when the
    service could not determine one. `type` is one of the eight wire types; see
    ocr-service/wire.py. Reading order is list order — never sort these.
    """
    id: str
    page: int
    type: str
    text: str
    bbox: Optional[List[float]] = None
    level: Optional[int] = None
    confidence: Optional[float] = None


@dataclass
class ParsedDocument:
    pages: List[PageContent]
    metadata: dict
    # Typed elements in reading order. Empty on the legacy (docling_enabled=False)
    # path, which is what makes the legacy chunker's fallback detectable.
    elements: List[Element] = field(default_factory=list)

    @property
    def text(self) -> str:
        return "\n".join(p.text for p in self.pages if p.text)


@dataclass
class ParentChunk:
    content: str
    page_start: int
    page_end: int
    source: str


@dataclass
class ChildChunk:
    content: str
    page: int
    source: str
    ocr_confidence: Optional[float] = None
    bbox: Optional[List[List[float]]] = None  # normalized rects covering this chunk
    # LLM-generated context situating this chunk in its document. None when
    # contextualization is disabled or failed; the chunk then embeds on content
    # alone, exactly as before this feature existed.
    context: Optional[str] = None
    # "table" for chunks derived from a table element, "text" otherwise. None on
    # the legacy path. Lets the eval harness score table questions separately.
    element_type: Optional[str] = None


_IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".webp"}
_MIME_BY_EXT = {
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".webp": "image/webp",
}


# --- Parsing ----------------------------------------------------------------

def _clamp01(v: float) -> float:
    return 0.0 if v < 0 else 1.0 if v > 1 else v


def _quad_to_norm_rect(
    quad: Optional[list], width: float, height: float
) -> Optional[List[float]]:
    """Convert an OCR polygon [[x,y], ...] (image-pixel coords) to a normalized
    axis-aligned rect [x0,y0,x1,y1] in [0,1]. Returns None if the quad or the
    page dimensions are missing."""
    if not quad or not width or not height:
        return None
    xs = [float(p[0]) for p in quad]
    ys = [float(p[1]) for p in quad]
    return [
        _clamp01(min(xs) / width), _clamp01(min(ys) / height),
        _clamp01(max(xs) / width), _clamp01(max(ys) / height),
    ]


def _ocr_page_image(
    image_bytes: bytes, filename: str
) -> Tuple[str, Optional[float], List[LayoutLine]]:
    """OCR a rendered page/image, degrading gracefully. Returns
    (text, avg_confidence, lines). Returns ("", None, []) when OCR is disabled or
    the service errors, so one bad page doesn't fail the whole document (the
    doc-level guard in the worker handles a fully-empty doc).

    `text` is "\\n".join of the non-blank line texts, matching the char layout the
    line bboxes are indexed against."""
    if not settings.ocr_enabled:
        return "", None, []
    try:
        body = ocr_image_lines(image_bytes, filename=filename)
    except OCRError:
        logger.warning("_ocr_page_image: OCR failed, treating page as empty file=%s", filename)
        return "", None, []

    width, height = body.get("width") or 0, body.get("height") or 0
    lines: List[LayoutLine] = []
    confs: List[float] = []
    for ln in body.get("lines", []):
        text = (ln.get("text") or "").strip()
        if not text:
            continue
        rect = _quad_to_norm_rect(ln.get("bbox"), width, height)
        lines.append(LayoutLine(text=text, bbox=rect or []))
        conf = ln.get("confidence")
        if isinstance(conf, (int, float)):
            confs.append(float(conf))

    text = "\n".join(l.text for l in lines)
    avg_conf = (sum(confs) / len(confs)) if confs else None
    return text, avg_conf, lines


def _native_pdf_lines(page) -> Tuple[List[LayoutLine], float, float]:
    """Extract per-line geometry from a native PDF page via PyMuPDF, normalized to
    [0,1] of the page rect. Returns (lines, page_width, page_height). Both PyMuPDF
    line bboxes and the pixmap/pdf.js viewport use a top-left origin, so no y-flip
    is needed."""
    rect = page.rect
    pw, ph = float(rect.width), float(rect.height)
    lines: List[LayoutLine] = []
    if not pw or not ph:
        return lines, pw, ph
    data = page.get_text("dict")
    for block in data.get("blocks", []):
        for line in block.get("lines", []):
            text = "".join(s.get("text", "") for s in line.get("spans", []))
            if not text.strip():
                continue
            x0, y0, x1, y1 = line.get("bbox", (0, 0, 0, 0))
            lines.append(LayoutLine(
                text=text,
                bbox=[_clamp01(x0 / pw), _clamp01(y0 / ph),
                      _clamp01(x1 / pw), _clamp01(y1 / ph)],
            ))
    return lines, pw, ph


def _parse_pdf(file_path: str, file_name: str) -> ParsedDocument:
    """Hybrid: use the native text layer when present; OCR scanned pages."""
    doc = fitz.open(file_path)
    pages: List[PageContent] = []
    ocr_pages = 0
    native_pages = 0
    for i, page in enumerate(doc):
        native_text = page.get_text() or ""
        if len(native_text.strip()) >= settings.ocr_min_text_chars:
            lines, pw, ph = _native_pdf_lines(page)
            # Rebuild text from the geometry lines so chunk char-offsets map back to
            # bboxes; fall back to raw text if geometry extraction found nothing.
            text = "\n".join(l.text for l in lines) if lines else native_text
            pages.append(PageContent(
                page=i + 1, text=text, source="native",
                lines=lines, width=pw, height=ph,
            ))
            native_pages += 1
        else:
            pix = page.get_pixmap(dpi=settings.ocr_dpi)
            img_bytes = pix.tobytes("png")
            text, conf, lines = _ocr_page_image(img_bytes, f"{file_name}#p{i + 1}.png")
            pages.append(PageContent(
                page=i + 1, text=text, source="ocr", ocr_confidence=conf,
                lines=lines, width=float(pix.width), height=float(pix.height),
            ))
            ocr_pages += 1
    metadata = {
        "page_count": len(pages),
        "ocr_pages": ocr_pages,
        "native_pages": native_pages,
        "ocr_engine": "paddleocr" if ocr_pages else None,
    }
    return ParsedDocument(pages=pages, metadata=metadata)


def _parse_docx(file_path: str) -> ParsedDocument:
    doc = DocxDocument(file_path)
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return ParsedDocument(
        pages=[PageContent(page=1, text=text, source="native")],
        metadata={"page_count": 1, "ocr_pages": 0, "native_pages": 1, "ocr_engine": None},
    )


def _parse_image(file_path: str, file_name: str) -> ParsedDocument:
    with open(file_path, "rb") as f:
        img_bytes = f.read()
    text, conf, lines = _ocr_page_image(img_bytes, file_name)
    return ParsedDocument(
        pages=[PageContent(page=1, text=text, source="ocr", ocr_confidence=conf, lines=lines)],
        metadata={"page_count": 1, "ocr_pages": 1, "native_pages": 0, "ocr_engine": "paddleocr"},
    )


def _elements_from_wire(body: dict) -> List[Element]:
    return [
        Element(
            id=el.get("id") or f"e{i}",
            page=int(el.get("page") or 1),
            type=el.get("type") or "paragraph",
            text=el.get("text") or "",
            bbox=el.get("bbox"),
            level=el.get("level"),
            confidence=el.get("confidence"),
        )
        for i, el in enumerate(body.get("elements") or [])
    ]


def _pages_from_wire(body: dict, elements: List[Element]) -> List[PageContent]:
    """Rebuild PageContent from the wire body.

    Chunking works off `elements`, but the contextualizer still needs
    `parsed.text` and per-page text, so every page gets its elements joined in
    reading order. `lines` stays empty — bbox attribution is per-element now.
    """
    text_by_page: dict = {}
    for el in elements:
        text_by_page.setdefault(el.page, []).append(el.text)

    pages: List[PageContent] = []
    for p in body.get("pages") or []:
        page_no = int(p.get("page") or 1)
        pages.append(PageContent(
            page=page_no,
            text="\n".join(text_by_page.get(page_no, [])),
            source=p.get("source") or "ocr",
            ocr_confidence=p.get("ocr_confidence"),
            width=p.get("width"),
            height=p.get("height"),
        ))

    # An element whose page number isn't covered by any emitted PageContent
    # would otherwise vanish from parsed.text/contextualization while still
    # being visible in `elements` — surface it rather than losing it silently.
    covered_pages = {p.page for p in pages}
    orphaned_pages = sorted(set(text_by_page) - covered_pages)
    if orphaned_pages:
        orphaned_elements = sum(len(text_by_page[pg]) for pg in orphaned_pages)
        logger.warning(
            "_pages_from_wire: %d element(s) reference page(s) %s absent from "
            "the wire body's pages list; their text is dropped from parsed.text",
            orphaned_elements, orphaned_pages,
        )
    return pages


def _parse_remote(file_path: str, file_name: str) -> ParsedDocument:
    """Structure-aware parse via ocr-service /parse. Handles PDF, DOCX and
    images uniformly — Docling detects the format itself.

    Errors propagate as OCRError/ParseTimeout: a failed parse must fail the
    document rather than silently degrade to structure-less chunks.
    """
    with open(file_path, "rb") as f:
        data = f.read()
    body = parse_document_remote(data, filename=file_name)
    elements = _elements_from_wire(body)
    pages = _pages_from_wire(body, elements)
    metadata = dict(body.get("metadata") or {})
    metadata["engine"] = "docling"
    return ParsedDocument(pages=pages, metadata=metadata, elements=elements)


def parse_document(file_path: str, file_name: str) -> ParsedDocument:
    """Parse a source file into typed elements + per-page text + metadata.

    With docling_enabled (the default) this is a single call to ocr-service
    /parse, which handles PDF, DOCX and images. The legacy per-format,
    line-based path is kept behind the flag for one release as the documented
    rollback; see the design spec.
    """
    ext = os.path.splitext(file_name)[1].lower()
    if settings.docling_enabled:
        parsed = _parse_remote(file_path, file_name)
    elif ext == ".pdf":
        parsed = _parse_pdf(file_path, file_name)
    elif ext == ".docx":
        parsed = _parse_docx(file_path)
    elif ext in _IMAGE_EXTS:
        parsed = _parse_image(file_path, file_name)
    else:
        parsed = ParsedDocument(pages=[], metadata={"page_count": 0})
    parsed.metadata["mime_type"] = _MIME_BY_EXT.get(ext)
    total_len = sum(len(p.text) for p in parsed.pages)
    logger.info(
        "parse_document: file=%s type=%s pages=%d elements=%d ocr_pages=%s text_len=%d",
        file_name, ext or "?", len(parsed.pages), len(parsed.elements),
        parsed.metadata.get("ocr_pages"), total_len,
    )
    return parsed


# --- Chunking ---------------------------------------------------------------

def _parent_splitter() -> RecursiveCharacterTextSplitter:
    return RecursiveCharacterTextSplitter.from_tiktoken_encoder(
        encoding_name="cl100k_base",
        chunk_size=1500,
        chunk_overlap=0,
    )


def _child_splitter() -> RecursiveCharacterTextSplitter:
    return RecursiveCharacterTextSplitter.from_tiktoken_encoder(
        encoding_name="cl100k_base",
        chunk_size=300,
        chunk_overlap=50,
    )


def chunk_text(text: str) -> Tuple[List[str], List[List[str]]]:
    """Low-level splitter for a single text block. Returns (parents,
    children_per_parent). children_per_parent[i] are the child chunks derived
    from parents[i]."""
    parents = _parent_splitter().split_text(text)
    child_splitter = _child_splitter()
    children_per_parent = [child_splitter.split_text(p) for p in parents]
    n_children = sum(len(c) for c in children_per_parent)
    logger.info(
        "chunk_text: input_len=%d parents=%d children=%d",
        len(text), len(parents), n_children,
    )
    return parents, children_per_parent


def _line_spans(lines: List[LayoutLine]) -> List[Tuple[int, int, List[float]]]:
    """Char [start, end) of each line within "\\n".join(l.text for l in lines),
    paired with its normalized rect. Mirrors how PageContent.text is built."""
    spans: List[Tuple[int, int, List[float]]] = []
    pos = 0
    for ln in lines:
        start = pos
        end = pos + len(ln.text)
        spans.append((start, end, ln.bbox))
        pos = end + 1  # account for the "\n" joiner
    return spans


def _rects_for_span(
    spans: List[Tuple[int, int, List[float]]], start: int, end: int
) -> List[List[float]]:
    """Rects of every line whose char span overlaps [start, end), in order."""
    return [box for s, e, box in spans if box and s < end and e > start]


def _find_from(haystack: str, needle: str, cursor: int) -> int:
    """str.find starting at cursor, falling back to a global find. Returns -1 if
    the substring is absent (e.g. splitter normalization changed whitespace)."""
    idx = haystack.find(needle, cursor)
    return idx if idx >= 0 else haystack.find(needle)


def _chunk_document_legacy(parsed: ParsedDocument) -> Tuple[List[ParentChunk], List[List[ChildChunk]]]:
    """Chunk page-by-page so every parent maps to exactly one page, giving exact
    page attribution. Children inherit their parent's page/source/confidence and,
    when the page has line geometry, the normalized bboxes of the lines they
    overlap (found by tracking each chunk's char offset within the page text)."""
    parents: List[ParentChunk] = []
    children_per_parent: List[List[ChildChunk]] = []
    for page in parsed.pages:
        if not page.text.strip():
            continue
        spans = _line_spans(page.lines)
        p_texts, c_per_p = chunk_text(page.text)
        p_cursor = 0
        for p_text, c_texts in zip(p_texts, c_per_p):
            p_start = _find_from(page.text, p_text, p_cursor)
            if p_start >= 0:
                p_cursor = p_start + len(p_text)
            parents.append(ParentChunk(
                content=p_text,
                page_start=page.page,
                page_end=page.page,
                source=page.source,
            ))
            child_objs: List[ChildChunk] = []
            c_cursor = 0
            for c in c_texts:
                c_off = _find_from(p_text, c, c_cursor)
                if c_off >= 0:
                    c_cursor = c_off + 1  # children overlap, so advance minimally
                if spans and p_start >= 0 and c_off >= 0:
                    g_start = p_start + c_off
                    rects = _rects_for_span(spans, g_start, g_start + len(c))
                else:
                    rects = []
                child_objs.append(ChildChunk(
                    content=c,
                    page=page.page,
                    source=page.source,
                    ocr_confidence=page.ocr_confidence,
                    bbox=rects,
                ))
            children_per_parent.append(child_objs)
    n_children = sum(len(c) for c in children_per_parent)
    logger.info(
        "chunk_document: pages=%d parents=%d children=%d",
        len(parsed.pages), len(parents), n_children,
    )
    return parents, children_per_parent


def chunk_document(parsed: ParsedDocument) -> Tuple[List[ParentChunk], List[List[ChildChunk]]]:
    """Chunk a parsed document.

    Uses layout-aware chunking when the parser returned typed elements; falls
    back to the legacy per-page line-based chunker when it did not (the
    docling_enabled=False rollback path).
    """
    if parsed.elements:
        from .chunking import chunk_elements  # local: chunking imports this module
        return chunk_elements(parsed.elements)
    return _chunk_document_legacy(parsed)


# --- Embeddings -------------------------------------------------------------

def _openai_client() -> OpenAI:
    # Embeddings are routed through OpenRouter's OpenAI-compatible /v1/embeddings.
    return OpenAI(
        api_key=settings.openrouter_api_key,
        base_url=settings.openrouter_base_url,
    )


def embed_text(text: str) -> List[float]:
    client = _openai_client()
    response = client.embeddings.create(
        model=settings.openai_embedding_model,
        input=[text],
        dimensions=settings.embedding_dim,
    )
    return response.data[0].embedding


def build_embedding_input(context: Optional[str], content: str) -> str:
    """Text actually sent to the embedding model for a child chunk.

    Contextual embeddings prepend the generated context so the vector carries
    document-level meaning. A chunk with no context embeds on content alone —
    this must stay byte-identical to the pre-contextual behaviour so disabling
    the feature is a true no-op.

    Format counterpart: app/services/rag/reranker.py:_rerank_text must produce
    the identical string, or retrieval and reranking see different text.
    """
    if not context:
        return content
    return f"{context}\n\n{content}"


def embed_chunks(chunks: List[str]) -> List[List[float]]:
    client = _openai_client()
    embeddings: List[List[float]] = []
    batch_size = 100
    total_batches = (len(chunks) + batch_size - 1) // batch_size
    logger.info("embed_chunks: total=%d batches=%d", len(chunks), total_batches)
    for i in range(0, len(chunks), batch_size):
        batch = chunks[i:i + batch_size]
        batch_num = i // batch_size + 1
        logger.debug("embed_chunks: batch=%d/%d size=%d", batch_num, total_batches, len(batch))
        response = client.embeddings.create(
            model=settings.openai_embedding_model,
            input=batch,
            dimensions=settings.embedding_dim,
        )
        embeddings.extend(item.embedding for item in response.data)
    logger.info("embed_chunks: done embeddings=%d", len(embeddings))
    return embeddings


# --- Persistence ------------------------------------------------------------

def store_chunks(
    db: Session,
    document_id: str,
    parents: List[ParentChunk],
    children_per_parent: List[List[ChildChunk]],
    child_embeddings: List[List[float]],
) -> None:
    """Insert parents then children. Children carry parent_id and page metadata.
    child_embeddings must be in the same flattened order as children_per_parent."""
    parent_rows = [
        DocumentParentChunk(
            document_id=document_id,
            parent_index=i,
            content=p.content,
            page_start=p.page_start,
            page_end=p.page_end,
            source=p.source,
        )
        for i, p in enumerate(parents)
    ]
    db.add_all(parent_rows)
    db.flush()  # populate parent_rows[*].id

    child_rows: List[DocumentChunk] = []
    embed_iter = iter(child_embeddings)
    global_idx = 0
    for parent_row, children in zip(parent_rows, children_per_parent):
        for child in children:
            child_rows.append(DocumentChunk(
                document_id=document_id,
                parent_id=parent_row.id,
                chunk_index=global_idx,
                content=child.content,
                embedding=next(embed_iter),
                page=child.page,
                source=child.source,
                ocr_confidence=child.ocr_confidence,
                bbox=child.bbox or None,
                context=child.context,
                element_type=child.element_type,
            ))
            global_idx += 1
    db.bulk_save_objects(child_rows)
    db.commit()
    logger.info(
        "store_chunks: parents=%d children=%d document_id=%s",
        len(parent_rows), len(child_rows), document_id,
    )
